"""
DOCX Parser using python-docx.
Extracts text and headings from Word documents.
"""
from docx import Document
from pathlib import Path


class DOCXParser:

    def parse(self, filepath: str) -> dict:
        doc = Document(filepath)

        full_text_parts = []
        headings = []
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings by style
            style_name = para.style.name.lower() if para.style else ""

            if "heading" in style_name or "title" in style_name:
                headings.append(text)
                current_section = text
                full_text_parts.append(f"\n## {text}\n")
            else:
                full_text_parts.append(text)

        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        full_text = "\n".join(full_text_parts)
        if table_texts:
            full_text += "\n\n" + "\n\n".join(table_texts)

        return {
            "full_text": full_text,
            "headings": headings,
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "title": Path(filepath).stem,
        }
